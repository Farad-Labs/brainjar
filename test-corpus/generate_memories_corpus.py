#!/usr/bin/env python3
"""
Generate a memories/ test corpus under test-corpus/ with date-rich files
across markdown, docx, pdf, and image formats.

The corpus models a real project's memory system:
  memories/semantic/   - extracted signals from conversations
  memories/episodic/   - daily memory captures
  memories/consolidated/ - deduplicated, refined, sorted

Dates are embedded in filenames, content headers, metadata, and EXIF.
"""

import os
import subprocess
import datetime
from pathlib import Path

# --- Python library imports ---
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from fpdf import FPDF

BASE = Path(__file__).parent / "docs" / "memories"

# macOS system font — used by ImageMagick for annotated images
FONT = "/System/Library/Fonts/Supplemental/Courier New Bold.ttf"

# ---------------------------------------------------------------------------
# Date palette — spread across a realistic range
# ---------------------------------------------------------------------------
DATES = [
    datetime.date(2025, 11, 3),
    datetime.date(2025, 11, 18),
    datetime.date(2025, 12, 1),
    datetime.date(2025, 12, 15),
    datetime.date(2025, 12, 22),
    datetime.date(2026, 1, 6),
    datetime.date(2026, 1, 14),
    datetime.date(2026, 1, 28),
    datetime.date(2026, 2, 5),
    datetime.date(2026, 2, 17),
    datetime.date(2026, 3, 1),
    datetime.date(2026, 3, 10),
    datetime.date(2026, 3, 15),
    datetime.date(2026, 3, 22),
    datetime.date(2026, 3, 28),
]

# Keep consistent with the existing corpus characters
PEOPLE = [
    "Sarah Chen", "Marcus Webb", "Priya Patel", "James Liu",
    "Diana Ross", "Elena Vasquez", "Alex Novak",
]


def ensure_dirs():
    for sub in ["semantic", "episodic", "consolidated"]:
        (BASE / sub).mkdir(parents=True, exist_ok=True)


# ===================================================================
# MARKDOWN FILES
# ===================================================================

SEMANTIC_MD = [
    {
        "filename": "user-intent-patterns_{date}.md",
        "title": "User Intent Patterns — {date_long}",
        "body": (
            "## Observed Patterns\n\n"
            "During the session on **{date_long}**, {person} demonstrated a preference for "
            "incremental refactoring over wholesale rewrites. Key signals:\n\n"
            "- Repeatedly asked for *small, testable changes*\n"
            "- Rejected a proposed bulk migration in favor of a phased approach\n"
            "- Referenced the Atlas incident on 2026-03-10 as justification\n\n"
            "## Confidence\n\n"
            "High — confirmed across 3 separate exchanges in the same session.\n\n"
            "## Related\n\n"
            "- See episodic entry for {date_iso}\n"
            "- Cross-references: incident-report.md, sprint-retro.md\n"
        ),
    },
    {
        "filename": "codebase-preference_{date}.md",
        "title": "Codebase Preferences — {date_long}",
        "body": (
            "## Extracted Signal\n\n"
            "As of **{date_long}**, the team strongly prefers:\n\n"
            "1. `Result<T, E>` over panics — zero `unwrap()` in production paths\n"
            "2. Trait-based abstractions for storage backends\n"
            "3. Integration tests over unit tests for database layers\n\n"
            "Source: conversation with {person} on {date_iso}.\n\n"
            "## Context\n\n"
            "This aligns with the architecture decisions documented in architecture.md "
            "and the onboarding guidance.\n"
        ),
    },
    {
        "filename": "search-relevance-feedback_{date}.md",
        "title": "Search Relevance Feedback — {date_long}",
        "body": (
            "## Feedback Captured\n\n"
            "On **{date_long}**, {person} noted that search results for "
            "\"Redis connection pooling\" were returning too many unrelated caching docs.\n\n"
            "### Expected Results\n"
            "- synonym-concepts.md (Redis caching layer)\n"
            "- incident-report.md (Redis OOM event on 2026-03-10)\n\n"
            "### Actual Results\n"
            "- Hidden-connections.md ranked #1 (irrelevant compliance doc)\n"
            "- Correct results ranked #4 and #7\n\n"
            "## Action\n\n"
            "Boost filename token matches. Consider date-decay weighting.\n"
        ),
    },
    {
        "filename": "domain-terminology_{date}.md",
        "title": "Domain Terminology Snapshot — {date_long}",
        "body": (
            "## Terms Learned\n\n"
            "Session date: **{date_long}**\n\n"
            "| Term | Meaning | First seen |\n"
            "|------|---------|------------|\n"
            "| AtlasQL | Domain-specific query language for the pipeline | 2025-11-03 |\n"
            "| Helix | Compliance governance framework | 2025-12-18 |\n"
            "| NATS migration | Move from Kafka to NATS for ingestion | 2026-01-14 |\n"
            "| Nyquist validation | Phase verification methodology | {date_iso} |\n\n"
            "Captured by: {person}\n"
        ),
    },
]

EPISODIC_MD = [
    {
        "filename": "{date}_daily-standup.md",
        "title": "Daily Standup — {date_long}",
        "body": (
            "**Date:** {date_iso}\n"
            "**Facilitator:** {person}\n\n"
            "## Summary\n\n"
            "- {person} is working on the NATS migration prototype\n"
            "- Blocked on ClickHouse schema approval (waiting since {prev_date})\n"
            "- PR #847 merged — Redis OOM fix from the March 10 incident\n\n"
            "## Action Items\n\n"
            "- [ ] {person}: Update ADR-003 by {next_date}\n"
            "- [ ] Review security findings from Alex Novak by end of week\n"
        ),
    },
    {
        "filename": "{date}_session-log.md",
        "title": "Session Log — {date_long}",
        "body": (
            "**Session date:** {date_iso}\n"
            "**Duration:** 45 minutes\n\n"
            "## Topics Discussed\n\n"
            "1. Refactoring the embedding pipeline — {person} wants to split "
            "the monolithic `process_document()` into smaller stages\n"
            "2. Reviewed test coverage: currently at 72%, target 85%\n"
            "3. Discussed temporal weighting for search results\n\n"
            "## Decisions\n\n"
            "- Use cosine similarity with a time-decay factor (half-life: 90 days from {date_iso})\n"
            "- Prioritize filename-date extraction in the next sprint\n\n"
            "## Follow-up\n\n"
            "- Schedule deep-dive on scoring algorithm for {next_date}\n"
        ),
    },
    {
        "filename": "{date}_debug-session.md",
        "title": "Debug Session Notes — {date_long}",
        "body": (
            "**Date:** {date_iso}\n"
            "**Participants:** {person}, Sarah Chen\n\n"
            "## Issue\n\n"
            "Search ranking returned stale documents above recent ones. "
            "A query for \"deployment checklist\" on {date_iso} surfaced a doc "
            "from 2025-11-03 over one from {prev_date}.\n\n"
            "## Root Cause\n\n"
            "No temporal signal in the scoring function. Pure cosine similarity "
            "treats all documents equally regardless of age.\n\n"
            "## Proposed Fix\n\n"
            "```\n"
            "score_final = score_semantic * (1.0 + temporal_boost(doc_date, query_date))\n"
            "```\n\n"
            "Target implementation: {next_date}\n"
        ),
    },
]

CONSOLIDATED_MD = [
    {
        "filename": "consolidated_{date}_weekly.md",
        "title": "Weekly Consolidation — Week of {date_long}",
        "body": (
            "**Period:** {date_iso} to {next_date}\n"
            "**Source entries:** 7 episodic, 3 semantic\n\n"
            "## Key Insights (Deduplicated)\n\n"
            "1. **Temporal scoring** is the top priority — mentioned in 4/7 daily entries\n"
            "2. **Redis stability** confirmed after the 2026-03-10 fix (PR #847)\n"
            "3. **ClickHouse schema** still pending approval (blocked since {prev_date})\n\n"
            "## Extracted Preferences\n\n"
            "- {person} prefers incremental changes (confirmed {date_iso})\n"
            "- Team consensus: integration tests > unit tests for DB layers\n\n"
            "## Stale Entries Removed\n\n"
            "- Duplicate standup notes from {prev_date} (identical content)\n"
            "- Superseded terminology entry (AtlasQL definition updated)\n"
        ),
    },
    {
        "filename": "consolidated_{date}_insights.md",
        "title": "Consolidated Insights — {date_long}",
        "body": (
            "**Consolidation date:** {date_iso}\n"
            "**Confidence threshold:** >= 0.7\n\n"
            "## High-Confidence Signals\n\n"
            "| Signal | Confidence | First Seen | Last Confirmed |\n"
            "|--------|-----------|------------|----------------|\n"
            "| Prefers small PRs | 0.95 | 2025-11-18 | {date_iso} |\n"
            "| No unwrap() in prod | 0.90 | 2025-12-01 | {date_iso} |\n"
            "| Time-decay scoring | 0.85 | 2026-02-05 | {date_iso} |\n"
            "| NATS over Kafka | 0.80 | 2026-01-14 | {prev_date} |\n\n"
            "## Deprecated Signals\n\n"
            "- \"Prefers mock-based tests\" — contradicted on {prev_date} by {person}\n"
        ),
    },
]


def fmt(template: str, date: datetime.date, person: str, idx: int) -> str:
    prev_date = date - datetime.timedelta(days=7)
    next_date = date + datetime.timedelta(days=7)
    return template.format(
        date=date.strftime("%Y-%m-%d"),
        date_iso=date.isoformat(),
        date_long=date.strftime("%B %d, %Y"),
        prev_date=prev_date.isoformat(),
        next_date=next_date.isoformat(),
        person=person,
    )


def write_markdown_files():
    # Semantic: 4 templates × ~3 dates each
    for i, d in enumerate(DATES[::4]):  # 4 dates
        for tmpl in SEMANTIC_MD:
            person = PEOPLE[i % len(PEOPLE)]
            fname = fmt(tmpl["filename"], d, person, i)
            title = fmt(tmpl["title"], d, person, i)
            body = fmt(tmpl["body"], d, person, i)
            path = BASE / "semantic" / fname
            path.write_text(f"# {title}\n\n{body}")
            print(f"  [md] {path.relative_to(BASE.parent)}")

    # Episodic: 3 templates × 5 dates
    for i, d in enumerate(DATES[::3]):  # 5 dates
        for tmpl in EPISODIC_MD:
            person = PEOPLE[i % len(PEOPLE)]
            fname = fmt(tmpl["filename"], d, person, i)
            title = fmt(tmpl["title"], d, person, i)
            body = fmt(tmpl["body"], d, person, i)
            path = BASE / "episodic" / fname
            path.write_text(f"# {title}\n\n{body}")
            print(f"  [md] {path.relative_to(BASE.parent)}")

    # Consolidated: 2 templates × 3 dates
    for i, d in enumerate(DATES[2::5]):  # 3 dates
        for tmpl in CONSOLIDATED_MD:
            person = PEOPLE[i % len(PEOPLE)]
            fname = fmt(tmpl["filename"], d, person, i)
            title = fmt(tmpl["title"], d, person, i)
            body = fmt(tmpl["body"], d, person, i)
            path = BASE / "consolidated" / fname
            path.write_text(f"# {title}\n\n{body}")
            print(f"  [md] {path.relative_to(BASE.parent)}")


# ===================================================================
# WORD DOCUMENTS (.docx)
# ===================================================================

DOCX_SPECS = [
    {
        "dir": "semantic",
        "filename": "architecture-preferences_{date}.docx",
        "title": "Architecture Preferences — {date_long}",
        "paragraphs": [
            "Document Date: {date_iso}",
            "Author: {person}",
            "",
            "The team has converged on a layered architecture with clear boundaries "
            "between ingestion, transformation, and storage. As discussed on {date_long}, "
            "the preference is for trait-based abstractions over concrete implementations.",
            "",
            "Key decisions recorded on {date_iso}:",
            "• Storage backend must implement the StorageProvider trait",
            "• All queries go through AtlasQL — no raw SQL in application code",
            "• Monitoring via OpenTelemetry (decided {prev_date})",
        ],
    },
    {
        "dir": "episodic",
        "filename": "{date}_incident-debrief.docx",
        "title": "Incident Debrief — {date_long}",
        "paragraphs": [
            "Incident Date: {date_iso}",
            "Severity: P2",
            "Lead: {person}",
            "",
            "On {date_long}, we experienced elevated error rates in the ingestion pipeline. "
            "The issue was traced to a misconfigured connection pool that was set too low "
            "for the current throughput (50K events/sec).",
            "",
            "Timeline:",
            "• {date_iso} 09:15 UTC — Alerts fired for error rate > 5%",
            "• {date_iso} 09:22 UTC — {person} acknowledged and began investigation",
            "• {date_iso} 09:45 UTC — Root cause identified: pool_size=10 vs required 50",
            "• {date_iso} 10:00 UTC — Fix deployed, error rate normalized",
            "",
            "Follow-up scheduled for {next_date}.",
        ],
    },
    {
        "dir": "consolidated",
        "filename": "consolidated_{date}_project-health.docx",
        "title": "Project Health Summary — {date_long}",
        "paragraphs": [
            "Reporting Period: {prev_date} to {date_iso}",
            "Prepared by: {person}",
            "",
            "Overall Status: GREEN",
            "",
            "Test coverage improved from 72% to 76% since {prev_date}. "
            "The NATS migration prototype was delivered on schedule. "
            "ClickHouse schema review is still pending (originally due {prev_date}).",
            "",
            "Risks:",
            "• Schema approval delay may impact Q2 2026 timeline",
            "• Security remediation from March 20 audit has 2 items remaining",
            "",
            "Next review: {next_date}",
        ],
    },
]


def write_docx_files():
    for spec in DOCX_SPECS:
        for i, d in enumerate(DATES[1::4]):  # 4 dates each
            person = PEOPLE[(i + 2) % len(PEOPLE)]
            fname = fmt(spec["filename"], d, person, i)
            title = fmt(spec["title"], d, person, i)

            doc = Document()

            # Set core metadata dates
            doc.core_properties.created = datetime.datetime(
                d.year, d.month, d.day, 10, 0, 0
            )
            doc.core_properties.modified = datetime.datetime(
                d.year, d.month, d.day, 16, 30, 0
            )
            doc.core_properties.author = person
            doc.core_properties.title = title
            doc.core_properties.subject = f"Memory entry for {d.isoformat()}"

            doc.add_heading(title, level=1)
            for para in spec["paragraphs"]:
                doc.add_paragraph(fmt(para, d, person, i))

            path = BASE / spec["dir"] / fname
            doc.save(str(path))
            print(f"  [docx] {path.relative_to(BASE.parent)}")


# ===================================================================
# PDF FILES
# ===================================================================

PDF_SPECS = [
    {
        "dir": "semantic",
        "filename": "scoring-analysis_{date}.pdf",
        "title": "Scoring Algorithm Analysis — {date_long}",
        "lines": [
            "Analysis Date: {date_iso}",
            "Analyst: {person}",
            "",
            "The current scoring function uses pure cosine similarity",
            "with no temporal component. Documents from {prev_date} and",
            "2025-11-03 receive identical ranking weight despite the",
            "90-day age difference.",
            "",
            "Proposed weighted score formula:",
            "  score = semantic_score * (1 + temporal_boost)",
            "  temporal_boost = exp(-age_days / half_life)",
            "",
            "With half_life=90 days from {date_iso}:",
            "  - 7-day-old doc: boost = 0.925",
            "  - 30-day-old doc: boost = 0.794",
            "  - 90-day-old doc: boost = 0.500",
            "  - 180-day-old doc: boost = 0.250",
            "",
            "Recommendation: implement before {next_date}.",
        ],
    },
    {
        "dir": "episodic",
        "filename": "{date}_search-quality-report.pdf",
        "title": "Search Quality Report — {date_long}",
        "lines": [
            "Report Date: {date_iso}",
            "Generated by: {person}",
            "",
            "Query: 'Redis connection issues'",
            "Expected top result: incident-report.md (2026-03-10)",
            "Actual top result: hidden-connections.md (irrelevant)",
            "",
            "Query: 'deployment checklist'",
            "Expected: onboarding.md",
            "Actual: roadmap.md (partial match only)",
            "",
            "Mean Reciprocal Rank: 0.42",
            "Target MRR: 0.75",
            "",
            "Temporal boost would improve MRR to estimated 0.68",
            "Filename token matching would add further improvement.",
            "",
            "Next measurement scheduled: {next_date}",
        ],
    },
    {
        "dir": "consolidated",
        "filename": "consolidated_{date}_metrics.pdf",
        "title": "Consolidated Metrics — {date_long}",
        "lines": [
            "Period: {prev_date} to {date_iso}",
            "Compiled by: {person}",
            "",
            "Search Quality Metrics:",
            "  MRR: 0.42 -> 0.58 (improving)",
            "  P@5: 0.60 -> 0.72",
            "  Temporal relevance score: 0.35 (target: 0.70)",
            "",
            "Memory System Metrics:",
            "  Semantic entries: 16",
            "  Episodic entries: 15",
            "  Consolidated entries: 6",
            "  Dedup ratio: 23% (7 duplicates removed)",
            "",
            "Infrastructure:",
            "  Embedding latency P95: 120ms",
            "  Index size: 2.4MB",
            "  Last full reindex: {prev_date}",
        ],
    },
]


def write_pdf_files():
    for spec in PDF_SPECS:
        for i, d in enumerate(DATES[2::3]):  # 5 dates each
            person = PEOPLE[(i + 1) % len(PEOPLE)]
            fname = fmt(spec["filename"], d, person, i)
            title = fmt(spec["title"], d, person, i)

            # fpdf2 core fonts are latin-1 only — strip unicode
            def latin1_safe(s: str) -> str:
                return s.replace("\u2014", "--").replace("\u2013", "-").replace("\u2022", "*")

            pdf = FPDF()
            pdf.add_page()

            # Metadata
            pdf.set_title(latin1_safe(title))
            pdf.set_author(person)
            pdf.set_subject(f"Memory entry for {d.isoformat()}")
            pdf.set_creation_date(datetime.datetime(d.year, d.month, d.day, 12, 0, 0))

            # Title
            pdf.set_font("Helvetica", "B", 16)
            pdf.cell(0, 12, latin1_safe(title), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(4)

            # Body
            pdf.set_font("Helvetica", "", 11)
            for line in spec["lines"]:
                text = latin1_safe(fmt(line, d, person, i))
                pdf.cell(0, 7, text, new_x="LMARGIN", new_y="NEXT")

            path = BASE / spec["dir"] / fname
            pdf.output(str(path))
            print(f"  [pdf] {path.relative_to(BASE.parent)}")


def write_scanned_pdf():
    """
    Create a flattened/scanned PDF — rendered as an image, no selectable text.
    This tests OCR extraction capabilities.
    """
    d = DATES[6]  # 2026-01-14
    person = "Priya Patel"
    fname = f"scanned-meeting-notes_{d.isoformat()}.pdf"
    title = f"Scanned Meeting Notes -- {d.strftime('%B %d, %Y')}"

    lines = [
        title,
        "",
        f"Date: {d.isoformat()}",
        f"Scanned by: {person}",
        "",
        "Attendees: Sarah Chen, Marcus Webb, Priya Patel, James Liu",
        "",
        "Agenda:",
        "1. Review NATS migration timeline (target: May 2026)",
        "2. ClickHouse schema approval -- blocked since 2026-01-06",
        "3. Search ranking improvements -- temporal weighting proposal",
        "",
        "Decisions:",
        "- Approved temporal decay with 90-day half-life",
        "- NATS prototype deadline extended to 2026-02-05",
        f"- Next review scheduled for {(d + datetime.timedelta(days=14)).isoformat()}",
        "",
        "Action items:",
        "- Marcus: draft ADR for scoring algorithm by 2026-01-21",
        "- Priya: benchmark embedding latency on new hardware",
        "- James: update client-notes.md with Meridian Q1 status",
    ]

    # Render text to image with ImageMagick
    text_block = "\n".join(lines)
    img_path = str(BASE / "episodic" / "scanned_tmp.png")
    subprocess.run(
        [
            "magick",
            "-size", "800x1000",
            "xc:white",
            "-gravity", "NorthWest",
            "-fill", "black",
            "-font", FONT,
            "-pointsize", "18",
            "-annotate", "+40+40", text_block,
            # Add slight noise to simulate scan artifacts
            "-attenuate", "0.15",
            "+noise", "Gaussian",
            "-quality", "85",
            img_path,
        ],
        check=True,
        capture_output=True,
    )

    # Convert image to PDF (flattened — no text layer)
    pdf_path = BASE / "episodic" / fname
    subprocess.run(
        [
            "magick",
            img_path,
            "-density", "150",
            str(pdf_path),
        ],
        check=True,
        capture_output=True,
    )

    os.remove(img_path)
    print(f"  [pdf/scan] {pdf_path.relative_to(BASE.parent)}")


# ===================================================================
# IMAGES (ImageMagick + EXIF via Python PIL)
# ===================================================================

IMAGE_SPECS = [
    {
        "dir": "semantic",
        "filename": "whiteboard-capture_{date}.jpg",
        "label": "Whiteboard: Scoring Algorithm\n{date_long}",
        "bg": "#2d3436",
        "fg": "#dfe6e9",
    },
    {
        "dir": "episodic",
        "filename": "{date}_architecture-diagram.jpg",
        "label": "Architecture Diagram\nCaptured {date_long}",
        "bg": "#0a3d62",
        "fg": "#f8c291",
    },
    {
        "dir": "episodic",
        "filename": "{date}_meeting-photo.jpg",
        "label": "Team Meeting\n{date_long}",
        "bg": "#6a0572",
        "fg": "#f0e5cf",
    },
    {
        "dir": "consolidated",
        "filename": "consolidated_{date}_dashboard.jpg",
        "label": "Metrics Dashboard\nWeek of {date_long}",
        "bg": "#1e3799",
        "fg": "#f5f6fa",
    },
]


def write_image_files():
    # Check if exiftool is available
    has_exiftool = subprocess.run(
        ["which", "exiftool"], capture_output=True
    ).returncode == 0

    if not has_exiftool:
        print("  (warning: exiftool not found — installing via brew)")
        subprocess.run(["brew", "install", "exiftool"], capture_output=True)
        has_exiftool = subprocess.run(
            ["which", "exiftool"], capture_output=True
        ).returncode == 0

    for spec in IMAGE_SPECS:
        for i, d in enumerate(DATES[::5]):  # 3 dates each
            person = PEOPLE[(i + 3) % len(PEOPLE)]
            fname = fmt(spec["filename"], d, person, i)
            label = fmt(spec["label"], d, person, i)

            path = BASE / spec["dir"] / fname
            exif_date = d.strftime("%Y:%m:%d 12:00:00")

            # Create image with ImageMagick
            subprocess.run(
                [
                    "magick",
                    "-size", "800x600",
                    f"xc:{spec['bg']}",
                    "-gravity", "Center",
                    "-fill", spec["fg"],
                    "-font", FONT,
                    "-pointsize", "36",
                    "-annotate", "+0+0", label,
                    "-gravity", "SouthEast",
                    "-pointsize", "14",
                    "-annotate", "+10+10", d.isoformat(),
                    str(path),
                ],
                check=True,
                capture_output=True,
            )

            # Set EXIF dates with exiftool
            if has_exiftool:
                subprocess.run(
                    [
                        "exiftool",
                        "-overwrite_original",
                        f"-DateTimeOriginal={exif_date}",
                        f"-CreateDate={exif_date}",
                        f"-ModifyDate={exif_date}",
                        f"-ImageDescription={label}",
                        f"-Artist={person}",
                        str(path),
                    ],
                    check=True,
                    capture_output=True,
                )
            else:
                print(f"    (no exiftool — EXIF dates not set for {fname})")

            print(f"  [img] {path.relative_to(BASE.parent)}")


# ===================================================================
# MAIN
# ===================================================================

def main():
    print("Creating memories/ corpus...\n")
    ensure_dirs()

    print("--- Markdown files ---")
    write_markdown_files()

    print("\n--- Word documents ---")
    write_docx_files()

    print("\n--- PDF files ---")
    write_pdf_files()

    print("\n--- Scanned/flattened PDF (OCR test) ---")
    write_scanned_pdf()

    print("\n--- Images ---")
    write_image_files()

    # Print summary
    total = 0
    for sub in ["semantic", "episodic", "consolidated"]:
        count = len(list((BASE / sub).iterdir()))
        total += count
        print(f"\n  {sub}/: {count} files")
    print(f"\n  Total: {total} files")
    print("\nDone!")


if __name__ == "__main__":
    main()
